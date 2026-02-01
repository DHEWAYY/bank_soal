import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- KONFIGURASI PENTING (JANGAN SALAH) ---
DOMAIN_UTAMA = "spacenet.my.id"
ADSENSE_ID = "ca-pub-8957427036950408"
ADS_TXT_CONTENT = "google.com, pub-8957427036950408, DIRECT, f08c47fec0942fa0"

# --- TEMPLATE HTML ---
TEMPLATE_HEADER = """<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{TITLE} | BankSoal.id</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <script async src="https://pagead2.googlesyndication.com/pagead/js/adsbygoogle.js?client={ADS_ID}" crossorigin="anonymous"></script>
</head>
<body class="bg-gray-50 text-gray-800 font-sans">
<nav class="bg-white border-b sticky top-0 z-50 shadow-sm"><div class="max-w-4xl mx-auto px-4 h-16 flex items-center justify-between"><a href="index.html" class="font-bold text-xl text-blue-600"><i class="fa-solid fa-book-open"></i> BankSoal.id</a></div></nav>
<main class="max-w-4xl mx-auto px-4 py-8">
<div class="w-full h-[100px] bg-gray-200 rounded-lg flex items-center justify-center mb-8 border-2 border-dashed border-gray-300 text-gray-400 font-bold text-xs">[IKLAN DISPLAY ADSENSE]</div>
"""

TEMPLATE_FOOTER = """
<div class="w-full h-[250px] bg-gray-200 rounded-lg flex items-center justify-center mt-8 border-2 border-dashed border-gray-300 text-gray-400 font-bold text-xs">[IKLAN KOTAK ADSENSE]</div>
</main><footer class="text-center py-8 text-gray-400 text-sm border-t mt-8 bg-white">&copy; 2026 BankSoal.id Engine</footer></body></html>"""

TEMPLATE_SOAL = """<article class="bg-white p-6 rounded-xl shadow-sm border mb-6"><div class="flex gap-3"><span class="bg-blue-100 text-blue-700 font-bold px-3 py-1 rounded h-fit text-sm">{NO}.</span><div class="w-full"><p class="text-lg font-medium mb-4">{PERTANYAAN}</p><div class="grid grid-cols-1 md:grid-cols-2 gap-3 mb-4"><div class="p-2 border rounded hover:bg-gray-50">A. {OPSI_A}</div><div class="p-2 border rounded hover:bg-gray-50">B. {OPSI_B}</div><div class="p-2 border rounded hover:bg-gray-50">C. {OPSI_C}</div><div class="p-2 border rounded hover:bg-gray-50">D. {OPSI_D}</div></div><details><summary class="cursor-pointer text-blue-600 font-semibold text-sm">Lihat Pembahasan</summary><div class="mt-2 p-3 bg-gray-50 rounded text-sm text-gray-700"><b>Jawaban: {JAWABAN}</b><br>{PEMBAHASAN}</div></details></div></div></article>"""

def create_docx(data, filename_base):
    try:
        doc = Document()
        meta = data.get('meta', {})
        doc.add_heading(meta.get('judul_bab', 'Latihan Soal'), 0)
        doc.add_paragraph(f"Mapel: {meta.get('mapel')} | Kelas: {meta.get('kelas')}")
        for q in data.get('soal_pg', []):
            doc.add_paragraph(f"{q['no']}. {q['tanya']}")
            doc.add_paragraph(f"A. {q['opsi_a']}  B. {q['opsi_b']}  C. {q['opsi_c']}  D. {q['opsi_d']}")
        path = f"docs/downloads/{filename_base}.docx"
        doc.save(path)
        return path.replace('docs/', '')
    except: return "#"

def generate():
    # 1. BIKIN STRUKTUR FOLDER BARU
    if not os.path.exists('docs'): os.makedirs('docs')
    if not os.path.exists('docs/downloads'): os.makedirs('docs/downloads')

    # 2. BIKIN FILE WAJIB (CNAME & ads.txt)
    with open('docs/CNAME', 'w') as f: f.write(DOMAIN_UTAMA)
    with open('docs/ads.txt', 'w') as f: f.write(ADS_TXT_CONTENT)
    
    # 3. BIKIN FILE .nojekyll (PENTING BIAR GAK ERROR DI GITHUB)
    with open('docs/.nojekyll', 'w') as f: f.write("")

    files = [f for f in os.listdir('data') if f.endswith('.json')]
    all_links = []

    print("🚀 Memproses Materi...")

    for filename in files:
        try:
            with open(f'data/{filename}', 'r', encoding='utf-8') as f: 
                data = json.load(f, strict=False) # STRICT FALSE BIAR GA ERROR JSON
            
            meta = data.get('meta', {})
            judul_bab = meta.get('judul_bab', 'Bank Soal')
            
            # Buat Konten HTML
            content_html = ""
            for q in data.get('soal_pg', []):
                content_html += TEMPLATE_SOAL.format(
                    NO=q['no'], PERTANYAAN=q['tanya'], 
                    OPSI_A=q['opsi_a'], OPSI_B=q['opsi_b'], OPSI_C=q['opsi_c'], OPSI_D=q['opsi_d'],
                    JAWABAN=q['jawaban'], PEMBAHASAN=q['pembahasan']
                )
            
            # Buat File Word
            nama_base = filename.replace('.json', '')
            link_docx = create_docx(data, nama_base)

            full_html = TEMPLATE_HEADER.format(TITLE=judul_bab, ADS_ID=ADSENSE_ID) + \
                        f"<div class='mb-6'><h1 class='text-2xl font-bold'>{judul_bab}</h1><p class='text-gray-500'>{meta.get('mapel')} - {meta.get('kelas')}</p><a href='{link_docx}' class='text-sm text-blue-600 hover:underline'><i class='fa-solid fa-download'></i> Download .DOCX</a></div>" + \
                        content_html + TEMPLATE_FOOTER
            
            out_name = f"{nama_base}.html"
            with open(f'docs/{out_name}', 'w', encoding='utf-8') as f: f.write(full_html)
            
            all_links.append(f'<a href="{out_name}" class="block p-5 bg-white border rounded-xl hover:shadow-lg transition hover:border-blue-400 group"><h3 class="font-bold text-gray-800 group-hover:text-blue-600 text-lg">{judul_bab}</h3><p class="text-sm text-gray-500 mt-2">{meta.get("mapel")} - {meta.get("kelas")}</p></a>')
            print(f"✅ Generated: {out_name}")

        except Exception as e:
            print(f"❌ Skip {filename} (Error: {e})")

    # Buat Index (Halaman Depan)
    grid_html = "<div class='grid grid-cols-1 md:grid-cols-2 gap-4'>" + "".join(all_links) + "</div>"
    index_html = TEMPLATE_HEADER.format(TITLE="Beranda", ADS_ID=ADSENSE_ID) + \
                 "<div class='text-center mb-12'><h1 class='text-4xl font-bold mb-4 text-gray-900'>Bank Soal Digital</h1><p class='text-lg text-gray-600'>Kumpulan latihan soal SD, SMP, SMA, SMK gratis tanpa login.</p></div>" + \
                 grid_html + TEMPLATE_FOOTER
    
    with open('docs/index.html', 'w', encoding='utf-8') as f: f.write(index_html)
    print("🏁 SELESAI! Folder 'docs' siap di-push.")

if __name__ == "__main__":
    generate()
